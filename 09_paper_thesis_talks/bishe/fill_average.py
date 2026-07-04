import docx

doc = docx.Document('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx')

table = doc.tables[0]

# Calculate average score
scores = [95, 96, 95, 94, 97, 95, 94, 96]
average_score = sum(scores) / len(scores)

# Put it in row 14, column 7
# (Column 7 corresponds to the score column, since columns 7 and 8 are merged in this row originally or something)
# The row 14 was: '平时成绩（每次评分的平均分）' in 0-6, empty in 7 and 8.
table.cell(14, 7).text = f"{average_score:.1f}"

# Also set font sizing uniformly for the text we added
from docx.shared import Pt
from docx.oxml.ns import qn

def set_font(cell):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11) # 五号

# Apply font to our inserted texts
for i in range(4, 12):
    set_font(table.cell(i, 0))
    set_font(table.cell(i, 2))
    set_font(table.cell(i, 7))

set_font(table.cell(0, 1))
set_font(table.cell(0, 4))
set_font(table.cell(0, 8))
set_font(table.cell(1, 1))
set_font(table.cell(1, 4))
set_font(table.cell(1, 8))
set_font(table.cell(14, 7))

doc.save('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx')
print("Successfully added average score and adjusted font!")
