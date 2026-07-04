import docx
from docx.shared import Pt
from docx.oxml.ns import qn

doc = docx.Document('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx')
table = doc.tables[0]

short_records = [
    "介绍课题背景，布置文献阅读，指导Mumax3软件基础。",
    "敲定任务书技术指标，明确三维Hopfion建模路线。",
    "审阅开题报告，指导STT/SOT代码及边界条件设置。",
    "检查稳态扫描，修正高阶交换系数公式并重算相图。",
    "分析自旋波驱动轨迹，指导设置吸收层避免反射干扰。",
    "讨论LIF神经元映射，指导拟合无外场漏电流参数。",
    "审查论文初稿，规范参考文献格式，优化图表排版。",
    "确认定稿与查重，预演答辩，指导竞争交换原理解释。"
]

def set_font(cell):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(11) # 五号

# Write shortened text to table rows 4 to 11
for i, short_text in enumerate(short_records):
    row_idx = 4 + i
    cell = table.cell(row_idx, 2)
    cell.text = short_text
    set_font(cell)

doc.save('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx')
print("Successfully shortened the content!")
