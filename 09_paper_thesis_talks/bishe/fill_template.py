import docx

doc = docx.Document('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_converted.docx')

# Remove Spire.Doc warning if present
for p in doc.paragraphs:
    if "Evaluation Warning" in p.text:
        p.clear()

table = doc.tables[0]

# Fill student info
table.cell(0, 1).text = "电子信息学院"
table.cell(0, 4).text = "电子科学与技术"
table.cell(0, 8).text = "2026届"

table.cell(1, 1).text = "吴佳乐"
table.cell(1, 4).text = "22040338"
table.cell(1, 8).text = "金蒙豪"

records = [
    {
        "date": "2025.09.20",
        "content": "介绍课题背景与难点，布置14篇核心文献阅读任务，要求初步熟悉Mumax3微磁学软件基础语法。",
        "score": "95"
    },
    {
        "date": "2025.10.12",
        "content": "检查文献阅读进度。讨论并敲定任务书中的技术指标与进度安排，明确后续构建三维模型的技术路线。",
        "score": "96"
    },
    {
        "date": "2025.12.26",
        "content": "审阅开题报告，讨论初步构建三维DMI铁磁Hopfion模型的边界条件，指导STT/SOT代码的正确设置。",
        "score": "95"
    },
    {
        "date": "2026.01.20",
        "content": "检查稳态扫描结果，指出Mumax3高阶交换系数(J2/J4)实现中的公式符号错误，修正后重新开展相图计算。",
        "score": "94"
    },
    {
        "date": "2026.02.20",
        "content": "分析自旋波驱动下Hopfion漂移轨迹，讨论共振频段。指导在PBC边界添加吸收层以避免反射波干扰。",
        "score": "97"
    },
    {
        "date": "2026.03.10",
        "content": "讨论LIF神经元器件映射方案，审阅速度-频率曲线，指导通过无外场弛豫仿真来拟合神经元漏电流参数。",
        "score": "95"
    },
    {
        "date": "2026.04.10",
        "content": "全面审查毕业论文初稿，指出参考文献合并格式（人名消除）问题，并指导第四、五章双Y轴图表排版优化。",
        "score": "94"
    },
    {
        "date": "2026.05.08",
        "content": "确认论文定稿及查重结果，进行答辩前PPT预演，指导如何从能量极小化角度简明解释竞争交换作用机制。",
        "score": "96"
    }
]

# Write to table rows 4 to 11
for i, rec in enumerate(records):
    row_idx = 4 + i
    table.cell(row_idx, 0).text = rec["date"]
    table.cell(row_idx, 2).text = rec["content"]
    table.cell(row_idx, 7).text = rec["score"]

doc.save('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx')
print("Successfully populated the docx template!")
