import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()
# Set default font
doc.styles['Normal'].font.name = '宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.styles['Normal'].font.size = Pt(12)

# Title
title = doc.add_paragraph('杭州电子科技大学\n毕业设计（论文）指导记录')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph() # Empty line

# Student info
info_p = doc.add_paragraph()
info_p.add_run('学    院：电子信息学院\t\t专    业：电子科学与技术\n')
info_p.add_run('班    级：22042011\t\t学生姓名：吴佳乐\n')
info_p.add_run('学    号：22040338\t\t指导教师：金蒙豪')
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph() # Empty line

records = [
    {
        "date": "2025年9月20日",
        "location": "办公室",
        "content": "导师详细介绍了磁性霍普夫子（Hopfion）的前沿研究背景及其在神经形态计算领域的应用潜力。讨论了本课题的难点，明确了使用 Mumax3 作为核心微磁学仿真工具。指导教师要求阅读指定的14篇核心文献，并布置了初步的软件熟悉任务和任务书的撰写工作。\n\n学生问题：Mumax3 对于三维复杂拓扑结构建模的局限性是什么？\n导师解答：建议首先从简单的二维斯格明子入手熟悉软件语法，再逐步通过自定义场项（如 Shifted 场）来实现三维竞争交换体系（如J1-J2-J4模型）的建模。"
    },
    {
        "date": "2025年12月26日",
        "location": "实验室",
        "content": "针对开题报告进行了第一次审稿。检查了文献综述部分，认为对磁性霍普夫子驱动机制的归纳还需完善。同时，检查了学生构建的初步三维 DMI 铁磁霍普夫子模型，指出了边界条件设置的注意事项。\n\n学生问题：如何在 Mumax3 中准确实现自旋转移力矩（STT）和自旋轨道力矩（SOT）的施加？\n导师解答：指导了 STT 和 SOT 项的具体代码书写方式，并强调需注意极化方向与电流密度的标度转换。建议开题报告中增加技术路线图以理清后续思路。"
    },
    {
        "date": "2026年1月20日",
        "location": "实验室",
        "content": "检查了基于竞争交换作用（Frustrated FM）的霍普夫子稳态参数扫描结果。发现之前 J4 和 J2 相互作用项由于 Mumax3 系数公式设置错误导致模型坍缩。重新推导了海森堡模型在微磁学连续极限下的等效系数表达式。\n\n学生问题：为什么之前加入高阶交换项（J2, J4）后，结构在 50ps 内会迅速坍塌？\n导师解答：指出交换系数转换公式中的符号错误，必须注意反铁磁耦合系数本身为负值。修正后指导学生重新开展各向异性（Ku）的临界值扫描。"
    },
    {
        "date": "2026年2月20日",
        "location": "实验室/线上",
        "content": "检查了自旋波（Spin Wave）驱动下 Hopfion 的运动行为。分析了微波场不同频率下的运动轨迹和漂移速度，发现特定频率区间内出现了共振加速现象。指导了数据后处理脚本的优化。\n\n学生问题：自旋波的激发源应该如何设置以保证平面波在传播过程中的平稳性？\n导师解答：建议采用局部交变磁场作为激发源，并在边界区域设置阻尼吸收层（absorbing boundary），防止自旋波在 PBC 边界处的反射对 Hopfion 造成干扰。"
    },
    {
        "date": "2026年3月10日",
        "location": "实验室",
        "content": "讨论了将 Hopfion 的非线性动力学特性映射到 LIF（Leaky Integrate-and-Fire）神经元模型的方案。审阅了学生提取的速度-频率曲线，指导学生如何利用这些曲线构建整合发放周期的数学模型。\n\n学生问题：Hopfion 在撤去驱动场后的阻尼弛豫过程（Leaky特性）该如何量化？\n导师解答：指导通过测量无外场状态下质心位置的回滞曲线或拓扑荷的演化来拟合漏电流时间常数，强调了理论计算与仿真结果必须自洽。"
    },
    {
        "date": "2026年4月10日",
        "location": "办公室",
        "content": "对毕业论文的第一稿进行了全面审查。指出了格式上的诸多问题，包括参考文献引用不规范（人名消除及合并引用问题），以及第四章、第五章图表排版不够紧凑。\n\n学生问题：动力学部分的图表信息量很大，如何排版才能更清晰地展示漂移轨迹与速度的关系？\n导师解答：建议将轨迹图与速度-时间曲线合并为双 Y 轴组合图，或者使用三维轨迹投影图。要求按照毕业设计撰写规范对格式进行逐项修改。"
    },
    {
        "date": "2026年5月8日",
        "location": "办公室",
        "content": "论文定稿及查重前最后一次确认。核心数据与结论已经完善，查重率符合学校规定。对即将到来的答辩报告 PPT 进行了初步审查，建议减少大段文字，增加仿真动画。\n\n学生问题：答辩时对于竞争交换模型的微观机制如果被评委提问，该如何简明扼要地回答？\n导师解答：建议从能量极小的角度解释，重点说明 J1-J2-J4 的竞争如何等效于一种长程的 DMI 作用，从而在无 DMI 材料中稳定拓扑结构。安排了模拟答辩时间。"
    }
]

for i, rec in enumerate(records):
    # Header row for each record
    p = doc.add_paragraph()
    p.add_run(f"第 {i+1} 次指导记录").bold = True
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    # Row 0
    table.cell(0,0).text = f"指导日期：{rec['date']}"
    table.cell(0,1).text = f"指导地点：{rec['location']}"
    
    # Row 1
    cell_content = table.cell(1,0)
    cell_content.merge(table.cell(1,1))
    p1 = cell_content.paragraphs[0]
    p1.add_run("指导内容及存在问题解答：\n").bold = True
    p1.add_run(rec['content'])
    
    # Row 2
    table.cell(2,0).text = "\n学生签名：\n"
    table.cell(2,1).text = "\n指导教师签名：\n"
    
    doc.add_paragraph() # Empty line between tables

output_path = "/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/毕业设计(论文)材料模板/04-指导记录_吴佳乐.docx"
doc.save(output_path)
print(f"Generated {output_path}")
