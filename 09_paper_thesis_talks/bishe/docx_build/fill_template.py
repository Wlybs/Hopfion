#!/usr/bin/env python3
"""
fill_template.py — 在模板 docx 原地替换文字内容，严格保留格式。

策略：
1. 以模板为底板（保留其 styles.xml, 页面设置, 页眉页脚, 段落格式）
2. 从 intermediate.docx（pandoc 输出）提取纯文本+图片+公式内容
3. 替换模板正文区域，复用模板的段落/Run 格式模板

关键原则：段落的 <w:pPr>（缩进/间距/对齐）和 <w:rPr>（字体/字号）全部保留模板的。
"""
import shutil
import zipfile
import os
import re
from pathlib import Path
from copy import deepcopy
from lxml import etree

WORKDIR = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/docx_build')
TPL_PATH = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/'
                '毕业设计(论文)材料模板/论文.docx')
INTER_PATH = WORKDIR / 'intermediate.docx'
OUT_PATH = WORKDIR / '论文_filled_v2.docx'

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def unpack(docx_path, dest):
    if dest.exists():
        shutil.rmtree(dest)
    dest.mkdir(parents=True)
    with zipfile.ZipFile(docx_path) as z:
        z.extractall(dest)


def repack(src_dir, out_path):
    if out_path.exists():
        out_path.unlink()
    with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(src_dir):
            for f in files:
                full = Path(root) / f
                z.write(full, full.relative_to(src_dir))


def parse_xml(path):
    return etree.parse(str(path))


def write_xml(tree, path):
    tree.write(str(path), xml_declaration=True, encoding='UTF-8', standalone=True)


def get_style(p_elem):
    """Get style name from a <w:p> element."""
    ppr = p_elem.find(f'{W}pPr')
    if ppr is not None:
        ps = ppr.find(f'{W}pStyle')
        if ps is not None:
            return ps.get(f'{W}val')
    return None


def get_text(p_elem):
    """Get all text from a paragraph element."""
    parts = []
    for t in p_elem.iter(f'{W}t'):
        if t.text:
            parts.append(t.text)
    return ''.join(parts)


def has_image(elem):
    """Check if element contains an image/drawing."""
    return len(elem.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')) > 0 or \
           len(elem.findall('.//' + '{urn:schemas-microsoft-com:vml}shape')) > 0 or \
           len(elem.findall('.//' + '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')) > 0


def extract_format_catalog(tpl_body):
    """Extract one representative <w:pPr> and <w:rPr> for each style from template."""
    catalog = {}  # style_name -> (pPr_element, rPr_element)
    for p in tpl_body.findall(f'{W}p'):
        style = get_style(p)
        if style and style not in catalog:
            ppr = p.find(f'{W}pPr')
            # Find first run's rPr
            rpr = None
            for r in p.findall(f'{W}r'):
                rpr = r.find(f'{W}rPr')
                if rpr is not None:
                    break
            catalog[style] = (deepcopy(ppr) if ppr is not None else None,
                              deepcopy(rpr) if rpr is not None else None)
    return catalog


def extract_body_format(tpl_body):
    """Extract the standard body text format (Normal with indent=304800)."""
    for p in tpl_body.findall(f'{W}p'):
        style = get_style(p)
        if style in (None, 'a', 'Normal', 'a0'):
            ppr = p.find(f'{W}pPr')
            if ppr is not None:
                ind = ppr.find(f'{W}ind')
                if ind is not None and ind.get(f'{W}firstLine'):
                    first = int(ind.get(f'{W}firstLine', '0'))
                    if first > 200000:  # ~2字符缩进
                        rpr = None
                        for r in p.findall(f'{W}r'):
                            rpr = r.find(f'{W}rPr')
                            if rpr is not None:
                                break
                        return deepcopy(ppr), deepcopy(rpr) if rpr else None
    return None, None


def make_paragraph(text, ppr_template, rpr_template):
    """Create a <w:p> with given format and text."""
    p = etree.SubElement(etree.Element('dummy'), f'{W}p')
    if ppr_template is not None:
        p.append(deepcopy(ppr_template))
    r = etree.SubElement(p, f'{W}r')
    if rpr_template is not None:
        r.append(deepcopy(rpr_template))
    t = etree.SubElement(r, f'{W}t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return p


def apply_template_rpr(run_elem, template_rpr):
    """Apply template's run properties to a run, keeping bold/italic/superscript."""
    if template_rpr is None:
        return
    existing_rpr = run_elem.find(f'{W}rPr')

    # Properties to KEEP from the original (content-specific formatting)
    keep_tags = {f'{W}b', f'{W}i', f'{W}vertAlign', f'{W}u',
                 f'{W}strike', f'{W}smallCaps'}
    kept = {}
    if existing_rpr is not None:
        for child in existing_rpr:
            if child.tag in keep_tags:
                kept[child.tag] = deepcopy(child)
        run_elem.remove(existing_rpr)

    new_rpr = deepcopy(template_rpr)
    # Remove from template the tags we'll keep from original
    for tag in keep_tags:
        for e in new_rpr.findall(tag):
            new_rpr.remove(e)
    # Add back kept properties
    for tag, elem in kept.items():
        new_rpr.append(elem)

    run_elem.insert(0, new_rpr)


def restyle_paragraph(p_elem, ppr_template, rpr_template):
    """Replace paragraph's formatting with template's, keep text and images."""
    # Replace pPr
    old_ppr = p_elem.find(f'{W}pPr')

    # Check if old pPr has sectPr (section break) — must preserve
    has_sect = False
    if old_ppr is not None:
        sect = old_ppr.find(f'{W}sectPr')
        if sect is not None:
            has_sect = True
            saved_sect = deepcopy(sect)
        p_elem.remove(old_ppr)

    if ppr_template is not None:
        new_ppr = deepcopy(ppr_template)
        if has_sect:
            new_ppr.append(saved_sect)
        p_elem.insert(0, new_ppr)

    # Apply rPr to each text run (not math/drawing runs)
    if rpr_template is not None:
        for r in p_elem.findall(f'{W}r'):
            # Skip runs that are purely drawing/image
            if r.find(f'{W}drawing') is not None or r.find(f'{W}pict') is not None:
                continue
            apply_template_rpr(r, rpr_template)


# ── Style mapping: pandoc style → template style + format ──

STYLE_MAP = {
    # pandoc → template style ID
    'Heading1': 'Heading1',
    'Heading2': 'Heading2',
    'Heading3': 'Heading3',
    '1': 'Heading1',
    '2': 'Heading2',
    '3': 'Heading3',
    'BodyText': 'BodyTextIndent2',  # pandoc body → template indented body
    'BodyTextIndent2': 'BodyTextIndent2',
    'FirstParagraph': None,  # use body format
    'Compact': None,
    'ImageCaption': None,   # will handle specially
    'TableCaption': None,
    'Caption': None,
}


def main():
    tpl_dir = WORKDIR / '_tpl_v2'
    inter_dir = WORKDIR / '_inter_v2'

    print("Step 1: Unpacking...")
    unpack(TPL_PATH, tpl_dir)
    unpack(INTER_PATH, inter_dir)

    print("Step 2: Parsing XML...")
    tpl_doc = parse_xml(tpl_dir / 'word' / 'document.xml')
    inter_doc = parse_xml(inter_dir / 'word' / 'document.xml')
    tpl_body = tpl_doc.getroot().find(f'{W}body')
    inter_body = inter_doc.getroot().find(f'{W}body')

    # Extract format catalog from template
    catalog = extract_format_catalog(tpl_body)
    body_ppr, body_rpr = extract_body_format(tpl_body)

    print(f"  Format catalog: {list(catalog.keys())}")
    print(f"  Body format found: pPr={'yes' if body_ppr is not None else 'no'}, "
          f"rPr={'yes' if body_rpr is not None else 'no'}")

    # ── Step 3: Identify template regions ──
    tpl_paras = tpl_body.findall(f'{W}p')
    tpl_all = list(tpl_body)

    # Find cover end (paragraph 18 = "年    月    日")
    cover_end_idx = 18

    # Find 摘要 title (paragraph 21)
    abstract_cn_title_idx = None
    abstract_en_title_idx = None
    toc_start_idx = None
    body_start_idx = None
    ack_start_idx = None
    ref_start_idx = None
    appendix_start_idx = None

    for i, p in enumerate(tpl_paras):
        txt = get_text(p)
        style = get_style(p)
        if '摘' in txt and '要' in txt and i < 30:
            abstract_cn_title_idx = i
        elif 'ABSTRACT' in txt and i < 50:
            abstract_en_title_idx = i
        elif '目' in txt and '录' in txt and i < 50:
            toc_start_idx = i
        elif style in ('Heading1', '1') and i > 50:
            if body_start_idx is None:
                body_start_idx = i
            txt_lower = txt.strip()
            if '致谢' in txt_lower:
                ack_start_idx = i
            elif '参考文献' in txt_lower:
                ref_start_idx = i
            elif '附录' in txt_lower:
                appendix_start_idx = i

    print(f"  CN abstract title: para {abstract_cn_title_idx}")
    print(f"  EN abstract title: para {abstract_en_title_idx}")
    print(f"  TOC start: para {toc_start_idx}")
    print(f"  Body start: para {body_start_idx}")
    print(f"  Acknowledgment: para {ack_start_idx}")
    print(f"  References: para {ref_start_idx}")
    print(f"  Appendix: para {appendix_start_idx}")

    # ── Step 4: Get intermediate content by section ──
    inter_paras = inter_body.findall(f'{W}p')

    # Parse intermediate into sections
    inter_sections = {
        'cn_abstract': [],   # paragraphs for CN abstract
        'cn_keywords': [],   # keywords paragraph
        'en_abstract': [],
        'en_keywords': [],
        'body': [],          # all body content elements (paras + tables)
        'ack': [],           # acknowledgment
        'refs': [],          # references
    }

    current = None
    for elem in inter_body:
        if elem.tag == f'{W}sectPr':
            continue
        if elem.tag == f'{W}p':
            txt = get_text(elem).strip()
            style = get_style(elem)

            if txt == '摘要' or txt == '摘    要':
                current = 'cn_abstract'
                continue
            elif txt == 'ABSTRACT':
                current = 'en_abstract'
                continue
            elif style in ('Heading1', '1'):
                if '致谢' in txt:
                    current = 'ack'
                    continue  # skip heading itself, template has it
                elif '参考文献' in txt:
                    current = 'refs'
                    continue
                else:
                    current = 'body'

            # Keywords detection
            if current == 'cn_abstract' and ('关键词' in txt or 'keywords' in txt.lower()):
                inter_sections['cn_keywords'].append(elem)
                current = None
                continue
            elif current == 'en_abstract' and ('key words' in txt.lower() or 'keywords' in txt.lower()):
                inter_sections['en_keywords'].append(elem)
                current = None
                continue

        if current:
            inter_sections[current].append(elem)

    for k, v in inter_sections.items():
        print(f"  Inter section '{k}': {len(v)} elements")

    # ── Step 5: Replace content in template ──
    print("\nStep 5: Replacing content...")

    # Helper: get format for a given style
    def get_format(style):
        if style in catalog:
            return catalog[style]
        # Default to body format
        return body_ppr, body_rpr

    # 5a: Replace CN abstract (paras between abstract title and keywords/ABSTRACT)
    # Template abstract: paras 22-24 are body, 26 is keywords
    # Delete template abstract body paras, insert new ones

    # Build a map of paragraph index → element
    para_idx_map = {}
    child_idx = 0
    for ci, child in enumerate(tpl_all):
        if child.tag == f'{W}p':
            para_idx_map[child_idx] = ci
            child_idx += 1

    def replace_region(start_para, end_para, new_elements, keep_first=False, keep_last=False):
        """Replace paragraphs [start_para..end_para) with new_elements.
        Restyle new elements to match template body format."""
        # Find child indices
        children = list(tpl_body)
        p_count = 0
        start_ci = None
        end_ci = None
        for ci, child in enumerate(children):
            if child.tag == f'{W}p':
                if p_count == start_para:
                    start_ci = ci
                if p_count == end_para:
                    end_ci = ci
                    break
                p_count += 1

        if start_ci is None:
            return
        if end_ci is None:
            end_ci = len(children)

        if keep_first:
            start_ci += 1
        if keep_last:
            end_ci -= 1

        # Remove old elements
        to_remove = children[start_ci:end_ci]
        for elem in to_remove:
            tpl_body.remove(elem)

        # Insert new elements at start_ci position
        insert_before = list(tpl_body)[start_ci] if start_ci < len(list(tpl_body)) else None

        for new_elem in new_elements:
            styled = deepcopy(new_elem)
            # Restyle based on its style
            style = get_style(styled)
            if style in ('Heading1', '1'):
                fmt = get_format('Heading1')
            elif style in ('Heading2', '2'):
                fmt = get_format('Heading2')
            elif style in ('Heading3', '3'):
                fmt = get_format('Heading3')
            else:
                fmt = (body_ppr, body_rpr)
                # Check if this is a caption (centered, short text about 图/表)
                txt = get_text(styled)
                if txt.startswith('图') or txt.startswith('表'):
                    if len(txt) < 100:
                        fmt = get_format('Normal') if 'Normal' not in catalog else (body_ppr, body_rpr)

            restyle_paragraph(styled, fmt[0], fmt[1])

            if insert_before is not None:
                tpl_body.insert(list(tpl_body).index(insert_before), styled)
            else:
                tpl_body.append(styled)

    # 5a: CN Abstract (keep title at para 21, replace body 22-25, keep keywords line at 26)
    cn_abs_body_start = abstract_cn_title_idx + 1  # 22
    # Find where CN abstract ends (before ABSTRACT or keywords)
    cn_abs_end = abstract_en_title_idx  # everything up to EN abstract title

    # Prepare new CN abstract paragraphs (restyled)
    cn_new = []
    for elem in inter_sections['cn_abstract']:
        cn_new.append(elem)
    for elem in inter_sections['cn_keywords']:
        cn_new.append(elem)

    replace_region(cn_abs_body_start, cn_abs_end, cn_new, keep_first=False, keep_last=False)

    # Recalculate indices after modification
    # Re-parse paragraph list
    tpl_paras = tpl_body.findall(f'{W}p')

    # Find EN abstract title again
    abstract_en_title_idx = None
    for i, p in enumerate(tpl_paras):
        if 'ABSTRACT' in get_text(p):
            abstract_en_title_idx = i
            break

    if abstract_en_title_idx is None:
        print("  WARNING: ABSTRACT title not found after CN abstract replacement!")
    else:
        print(f"  EN abstract now at para {abstract_en_title_idx}")

    # 5b: EN Abstract
    # Find EN keywords line and TOC start after it
    en_abs_start = abstract_en_title_idx + 1
    # Find next structural element
    en_abs_end = None
    for i in range(en_abs_start, len(tpl_paras)):
        txt = get_text(tpl_paras[i])
        if '目' in txt and '录' in txt:
            en_abs_end = i
            break

    if en_abs_end:
        en_new = []
        for elem in inter_sections['en_abstract']:
            en_new.append(elem)
        for elem in inter_sections['en_keywords']:
            en_new.append(elem)
        replace_region(en_abs_start, en_abs_end, en_new)
        print(f"  EN abstract replaced ({len(en_new)} paras)")

    # 5c: Recalculate and find body region
    tpl_paras = tpl_body.findall(f'{W}p')
    body_start = None
    ack_idx = None
    ref_idx = None
    app_idx = None
    for i, p in enumerate(tpl_paras):
        style = get_style(p)
        txt = get_text(p).strip()
        if style in ('Heading1', '1') and body_start is None and i > 30:
            body_start = i
        if style in ('Heading1', '1') and '致谢' in txt:
            ack_idx = i
        if style in ('Heading1', '1') and '参考文献' in txt:
            ref_idx = i
        if style in ('Heading1', '1') and '附录' in txt:
            app_idx = i

    print(f"\n  Body start: {body_start}, 致谢: {ack_idx}, 参考文献: {ref_idx}, 附录: {app_idx}")

    # 5d: Replace main body (from first Heading1 to 致谢)
    if body_start is not None and ack_idx is not None:
        replace_region(body_start, ack_idx, inter_sections['body'])
        print(f"  Body replaced ({len(inter_sections['body'])} elements)")

    # Recalculate
    tpl_paras = tpl_body.findall(f'{W}p')
    ack_idx = ref_idx = app_idx = None
    for i, p in enumerate(tpl_paras):
        style = get_style(p)
        txt = get_text(p).strip()
        if style in ('Heading1', '1'):
            if '致谢' in txt:
                ack_idx = i
            elif '参考文献' in txt:
                ref_idx = i
            elif '附录' in txt:
                app_idx = i

    # 5e: Replace 致谢 body (between 致谢 heading and 参考文献 heading)
    if ack_idx is not None and ref_idx is not None:
        replace_region(ack_idx + 1, ref_idx, inter_sections['ack'])
        print(f"  致谢 replaced ({len(inter_sections['ack'])} paras)")

    # Recalculate
    tpl_paras = tpl_body.findall(f'{W}p')
    ref_idx = app_idx = None
    for i, p in enumerate(tpl_paras):
        style = get_style(p)
        txt = get_text(p).strip()
        if style in ('Heading1', '1'):
            if '参考文献' in txt:
                ref_idx = i
            elif '附录' in txt:
                app_idx = i

    # 5f: Replace references (between 参考文献 heading and 附录/end)
    if ref_idx is not None:
        ref_end = app_idx if app_idx is not None else len(tpl_paras)
        replace_region(ref_idx + 1, ref_end, inter_sections['refs'])
        print(f"  参考文献 replaced ({len(inter_sections['refs'])} paras)")

    # ── Step 6: Copy images from intermediate ──
    print("\nStep 6: Merging images...")
    inter_media = inter_dir / 'word' / 'media'
    tpl_media = tpl_dir / 'word' / 'media'
    tpl_media.mkdir(exist_ok=True)

    if inter_media.exists():
        for f in inter_media.iterdir():
            dest = tpl_media / f.name
            if not dest.exists():
                shutil.copy2(f, dest)

    # Merge image relationships from intermediate
    inter_rels = parse_xml(inter_dir / 'word' / '_rels' / 'document.xml.rels')
    tpl_rels = parse_xml(tpl_dir / 'word' / '_rels' / 'document.xml.rels')
    ir = inter_rels.getroot()
    tr = tpl_rels.getroot()

    existing_ids = {r.get('Id') for r in tr}
    existing_types = {r.get('Type') for r in tr}
    SINGLETONS = {'styles', 'settings', 'webSettings', 'fontTable',
                  'numbering', 'footnotes', 'endnotes', 'theme'}

    for rel in ir:
        rid = rel.get('Id')
        target = rel.get('Target', '')
        reltype = rel.get('Type', '')
        type_base = reltype.rsplit('/', 1)[-1] if reltype else ''

        if rid in existing_ids:
            continue
        if type_base in SINGLETONS and reltype in existing_types:
            continue

        target_path = tpl_dir / 'word' / target
        if target.startswith('http') or target.startswith('#') or target_path.exists():
            tr.append(deepcopy(rel))

    write_xml(tpl_rels, tpl_dir / 'word' / '_rels' / 'document.xml.rels')

    # Merge Content_Types
    inter_ct = parse_xml(inter_dir / '[Content_Types].xml')
    tpl_ct = parse_xml(tpl_dir / '[Content_Types].xml')
    ir_ct = inter_ct.getroot()
    tr_ct = tpl_ct.getroot()

    existing_exts = {e.get('Extension') for e in tr_ct if e.get('Extension')}
    for elem in ir_ct:
        ext = elem.get('Extension')
        if ext and ext not in existing_exts:
            tr_ct.append(deepcopy(elem))
            existing_exts.add(ext)

    write_xml(tpl_ct, tpl_dir / '[Content_Types].xml')

    # ── Step 7: Write and repack ──
    print("\nStep 7: Writing and repacking...")
    write_xml(tpl_doc, tpl_dir / 'word' / 'document.xml')
    repack(tpl_dir, OUT_PATH)

    print(f"\nDONE: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024 / 1024:.1f} MB")

    # Cleanup
    shutil.rmtree(tpl_dir, ignore_errors=True)
    shutil.rmtree(inter_dir, ignore_errors=True)

    # Verify
    from docx import Document
    doc = Document(str(OUT_PATH))
    print(f"\nVerification: {len(doc.paragraphs)} paras, {len(doc.tables)} tables")
    img_count = len([r for r in doc.part.rels.values() if 'image' in r.reltype])
    print(f"Images: {img_count}")


if __name__ == '__main__':
    main()
