#!/usr/bin/env python3
"""
restyle.py — 以 intermediate.docx 为内容底，强制覆盖为模板格式。

工作流：
1. 从 论文_filled.docx（v1，已有封面+内容）出发
2. 替换 styles.xml 为模板的（字体/间距定义源）
3. 逐段按角色重刷 pPr/rPr（体格式、图注格式、标题格式等）
4. 复制模板的页面设置（sectPr）
"""
import shutil, zipfile, os
from pathlib import Path
from copy import deepcopy
from lxml import etree

WORKDIR = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/docx_build')
V1_PATH = WORKDIR / '论文_filled.docx'  # v1 output (has cover + all content)
TPL_PATH = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/'
                '毕业设计(论文)材料模板/论文.docx')
OUT_PATH = WORKDIR / '论文_filled_v2.docx'

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def unpack(path, dest):
    if dest.exists(): shutil.rmtree(dest)
    dest.mkdir(parents=True)
    with zipfile.ZipFile(path) as z: z.extractall(dest)

def repack(src, out):
    if out.exists(): out.unlink()
    with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(src):
            for f in files:
                full = Path(root) / f
                z.write(full, full.relative_to(src))

def get_style(p):
    ppr = p.find(f'{W}pPr')
    if ppr is not None:
        ps = ppr.find(f'{W}pStyle')
        if ps is not None:
            return ps.get(f'{W}val')
    return None

def get_text(p):
    return ''.join(t.text or '' for t in p.iter(f'{W}t'))

def has_drawing(p):
    ns_wp = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    ns_v = '{urn:schemas-microsoft-com:vml}'
    return bool(p.findall(f'.//{ns_wp}inline') or p.findall(f'.//{ns_wp}anchor') or p.findall(f'.//{ns_v}shape'))


# ── Format builders ──

def make_body_ppr():
    """标准正文段落格式: 首行缩进2字符, 行距20pt exact"""
    ppr = etree.Element(f'{W}pPr')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}firstLine', '480')
    ind.set(f'{W}firstLineChars', '200')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', '400')
    spc.set(f'{W}lineRule', 'exact')
    # rPr inside pPr (paragraph-level default run props)
    rpr = etree.SubElement(ppr, f'{W}rPr')
    rf = etree.SubElement(rpr, f'{W}rFonts')
    rf.set(f'{W}ascii', '宋体')
    rf.set(f'{W}hAnsi', '宋体')
    rf.set(f'{W}hint', 'eastAsia')
    sz = etree.SubElement(rpr, f'{W}sz')
    sz.set(f'{W}val', '24')
    return ppr

def make_body_rpr():
    """标准正文 Run 格式: 宋体 小四(12pt=24半磅)"""
    rpr = etree.Element(f'{W}rPr')
    rf = etree.SubElement(rpr, f'{W}rFonts')
    rf.set(f'{W}ascii', '宋体')
    rf.set(f'{W}hAnsi', '宋体')
    rf.set(f'{W}hint', 'eastAsia')
    sz = etree.SubElement(rpr, f'{W}sz')
    sz.set(f'{W}val', '24')
    return rpr

def make_caption_ppr():
    """图注/表注格式: 居中, 行距20pt, 五号"""
    ppr = etree.Element(f'{W}pPr')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', '400')
    spc.set(f'{W}lineRule', 'exact')
    jc = etree.SubElement(ppr, f'{W}jc')
    jc.set(f'{W}val', 'center')
    rpr = etree.SubElement(ppr, f'{W}rPr')
    szcs = etree.SubElement(rpr, f'{W}szCs')
    szcs.set(f'{W}val', '21')
    return ppr

def make_caption_rpr():
    rpr = etree.Element(f'{W}rPr')
    szcs = etree.SubElement(rpr, f'{W}szCs')
    szcs.set(f'{W}val', '21')
    return rpr

def make_keywords_ppr(is_cn=True):
    """关键词格式: 黑体, 左缩进"""
    ppr = etree.Element(f'{W}pPr')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}left', '840')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', '400')
    spc.set(f'{W}lineRule', 'exact')
    rpr = etree.SubElement(ppr, f'{W}rPr')
    rf = etree.SubElement(rpr, f'{W}rFonts')
    rf.set(f'{W}eastAsia', '黑体')
    sz = etree.SubElement(rpr, f'{W}sz')
    sz.set(f'{W}val', '24')
    return ppr

def make_abstract_body_ppr():
    """摘要正文: 楷体, 行距20pt"""
    ppr = etree.Element(f'{W}pPr')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', '400')
    spc.set(f'{W}lineRule', 'exact')
    rpr = etree.SubElement(ppr, f'{W}rPr')
    sz = etree.SubElement(rpr, f'{W}sz')
    sz.set(f'{W}val', '24')
    return ppr

def make_abstract_body_rpr():
    rpr = etree.Element(f'{W}rPr')
    rf = etree.SubElement(rpr, f'{W}rFonts')
    rf.set(f'{W}ascii', '楷体_GB2312')
    rf.set(f'{W}eastAsia', '楷体_GB2312')
    sz = etree.SubElement(rpr, f'{W}sz')
    sz.set(f'{W}val', '24')
    return rpr

def make_ref_ppr():
    """参考文献条目: 无缩进, 行距20pt, 宋体五号"""
    ppr = etree.Element(f'{W}pPr')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', '400')
    spc.set(f'{W}lineRule', 'exact')
    return ppr


# ── Core restyle functions ──

def restyle_ppr(p, new_ppr):
    """Replace paragraph's pPr, preserving sectPr if present."""
    old = p.find(f'{W}pPr')
    sect = None
    if old is not None:
        sect = old.find(f'{W}sectPr')
        if sect is not None:
            sect = deepcopy(sect)
        p.remove(old)

    fresh = deepcopy(new_ppr)
    if sect is not None:
        fresh.append(sect)
    p.insert(0, fresh)

def restyle_runs(p, new_rpr):
    """Apply rPr template to all text runs, keeping bold/italic/sup."""
    KEEP = {f'{W}b', f'{W}bCs', f'{W}i', f'{W}iCs',
            f'{W}vertAlign', f'{W}u', f'{W}strike', f'{W}smallCaps'}

    for r in p.findall(f'{W}r'):
        # Skip image/drawing runs
        if r.find(f'{W}drawing') is not None or r.find(f'{W}pict') is not None:
            continue

        old_rpr = r.find(f'{W}rPr')
        kept = {}
        if old_rpr is not None:
            for child in old_rpr:
                if child.tag in KEEP:
                    kept[child.tag] = deepcopy(child)
            r.remove(old_rpr)

        fresh = deepcopy(new_rpr)
        for tag, elem in kept.items():
            # Remove conflicting from template
            for e in fresh.findall(tag):
                fresh.remove(e)
            fresh.append(elem)
        r.insert(0, fresh)


def classify_paragraph(p, context):
    """Determine paragraph role: heading/body/caption/keywords/ref/abstract."""
    style = get_style(p)
    txt = get_text(p).strip()

    if style in ('1',):
        return 'heading1'
    if style in ('2',):
        return 'heading2'
    if style in ('3', 'Heading3', 'Heading4'):
        return 'heading3'

    if context == 'cn_abstract':
        if '关键词' in txt:
            return 'cn_keywords'
        return 'cn_abstract_body'
    if context == 'en_abstract':
        if 'key word' in txt.lower() or 'keywords' in txt.lower():
            return 'en_keywords'
        return 'en_abstract_body'
    if context == 'refs':
        return 'reference'

    # Body text heuristics
    if has_drawing(p):
        return 'figure'  # paragraph containing an image
    if txt and (txt.startswith('图') or txt.startswith('表')) and len(txt) < 120:
        return 'caption'

    return 'body'


def main():
    v1_dir = WORKDIR / '_v1_unpack'
    tpl_dir = WORKDIR / '_tpl_restyle'

    print("Step 1: Unpacking...")
    unpack(V1_PATH, v1_dir)
    unpack(TPL_PATH, tpl_dir)

    # Step 2: Copy template's styles.xml and settings
    print("Step 2: Overwriting styles with template's...")
    for fname in ['styles.xml']:
        src = tpl_dir / 'word' / fname
        dst = v1_dir / 'word' / fname
        if src.exists():
            shutil.copy2(src, dst)
            print(f"  Copied {fname}")

    # Step 3: Copy template's section properties (margins, headers, footers)
    print("Step 3: Copying section properties from template...")
    tpl_doc = etree.parse(str(tpl_dir / 'word' / 'document.xml'))
    v1_doc = etree.parse(str(v1_dir / 'word' / 'document.xml'))
    tpl_body = tpl_doc.getroot().find(f'{W}body')
    v1_body = v1_doc.getroot().find(f'{W}body')

    # Replace final sectPr (document-level section properties)
    tpl_sect = tpl_body.find(f'{W}sectPr')
    v1_sect = v1_body.find(f'{W}sectPr')
    if tpl_sect is not None and v1_sect is not None:
        idx = list(v1_body).index(v1_sect)
        v1_body.remove(v1_sect)
        v1_body.insert(idx, deepcopy(tpl_sect))
        print("  Replaced document sectPr")

    # Copy header/footer files from template
    for subdir in ['']:
        for pattern in ['header*.xml', 'footer*.xml']:
            import glob
            for src in (tpl_dir / 'word').glob(pattern):
                dst = v1_dir / 'word' / src.name
                shutil.copy2(src, dst)
                print(f"  Copied {src.name}")

    # Step 4: Restyle all paragraphs
    print("\nStep 4: Restyling paragraphs...")
    paras = v1_body.findall(f'{W}p')

    # Detect regions
    context = 'cover'  # start in cover
    stats = {}

    for i, p in enumerate(paras):
        style = get_style(p)
        txt = get_text(p).strip()

        # Context transitions
        if '年' in txt and '月' in txt and '日' in txt and i < 25:
            context = 'post_cover'
            continue  # don't restyle cover
        if i < 20:
            continue  # cover area, skip

        if txt in ('摘要', '摘    要') or (style == 'a3' and '摘' in txt and '要' in txt):
            context = 'cn_abstract'
            continue  # keep the 摘要 title as-is (template's)
        if txt == 'ABSTRACT' or (style == 'a3' and 'ABSTRACT' in txt):
            context = 'en_abstract'
            continue
        if '目' in txt and '录' in txt and i < 80:
            context = 'toc'
            continue
        if style in ('10', '22'):  # TOC styles
            continue  # skip TOC entries
        # Transition to body: first Heading1 after abstract/TOC
        if style in ('1',) and context in ('toc', 'en_abstract', 'post_cover', 'cn_abstract'):
            context = 'body'
        # Also: after en_keywords, next non-empty paragraph starts body
        if context == 'en_abstract' and '关键词' not in txt and 'key' not in txt.lower():
            # Check if we passed keywords already (heuristic: Heading style = body start)
            if style in ('1', '2', '3'):
                context = 'body'
        if style in ('1',) and '致谢' in txt:
            context = 'ack'
            continue  # keep heading
        if style in ('1',) and '参考文献' in txt:
            context = 'refs'
            continue
        if style in ('1',) and '附录' in txt:
            context = 'appendix'
            continue

        if context in ('cover', 'post_cover', 'toc', 'appendix'):
            continue

        # For ack context, classify as body (same format as body text)
        if context == 'ack':
            role = classify_paragraph(p, 'body')
        else:
            role = classify_paragraph(p, context)
        stats[role] = stats.get(role, 0) + 1

        # Apply formatting based on role
        if role == 'heading1':
            # Keep heading style, just ensure line spacing
            ppr = p.find(f'{W}pPr')
            if ppr is not None:
                spc = ppr.find(f'{W}spacing')
                if spc is None:
                    spc = etree.SubElement(ppr, f'{W}spacing')
                spc.set(f'{W}line', '400')
                spc.set(f'{W}lineRule', 'exact')
        elif role == 'heading2':
            ppr = p.find(f'{W}pPr')
            if ppr is not None:
                spc = ppr.find(f'{W}spacing')
                if spc is None:
                    spc = etree.SubElement(ppr, f'{W}spacing')
                spc.set(f'{W}line', '400')
                spc.set(f'{W}lineRule', 'exact')
        elif role == 'heading3':
            ppr = p.find(f'{W}pPr')
            if ppr is not None:
                spc = ppr.find(f'{W}spacing')
                if spc is None:
                    spc = etree.SubElement(ppr, f'{W}spacing')
                spc.set(f'{W}line', '400')
                spc.set(f'{W}lineRule', 'exact')
        elif role == 'body':
            restyle_ppr(p, make_body_ppr())
            restyle_runs(p, make_body_rpr())
        elif role == 'cn_abstract_body':
            restyle_ppr(p, make_abstract_body_ppr())
            restyle_runs(p, make_abstract_body_rpr())
        elif role == 'en_abstract_body':
            restyle_ppr(p, make_body_ppr())  # EN abstract uses 宋体 same as body
            restyle_runs(p, make_body_rpr())
        elif role in ('cn_keywords', 'en_keywords'):
            restyle_ppr(p, make_keywords_ppr())
        elif role == 'caption':
            restyle_ppr(p, make_caption_ppr())
            restyle_runs(p, make_caption_rpr())
        elif role == 'figure':
            # Image paragraph: center, no indent
            restyle_ppr(p, make_caption_ppr())
        elif role == 'reference':
            restyle_ppr(p, make_ref_ppr())
            restyle_runs(p, make_body_rpr())
        elif role == 'ack':
            restyle_ppr(p, make_body_ppr())
            restyle_runs(p, make_body_rpr())

    print(f"  Role stats: {stats}")

    # Step 5: Write and repack
    print("\nStep 5: Repacking...")
    etree.ElementTree(v1_doc.getroot()).write(
        str(v1_dir / 'word' / 'document.xml'),
        xml_declaration=True, encoding='UTF-8', standalone=True)

    repack(v1_dir, OUT_PATH)
    print(f"\nDONE: {OUT_PATH}")
    print(f"Size: {OUT_PATH.stat().st_size / 1024 / 1024:.1f} MB")

    # Copy to target (skip if locked by Word)
    final = TPL_PATH.parent / '论文_filled.docx'
    try:
        shutil.copy2(OUT_PATH, final)
        print(f"Copied to: {final}")
    except PermissionError:
        print(f"  WARN: {final} locked, output at: {OUT_PATH}")

    # Verify
    from docx import Document
    doc = Document(str(OUT_PATH))
    print(f"Verification: {len(doc.paragraphs)} paras, {len(doc.tables)} tables")

    # Cleanup
    shutil.rmtree(v1_dir, ignore_errors=True)
    shutil.rmtree(tpl_dir, ignore_errors=True)


if __name__ == '__main__':
    main()
