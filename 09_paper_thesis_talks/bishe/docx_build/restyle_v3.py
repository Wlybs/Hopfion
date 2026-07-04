#!/usr/bin/env python3
"""
restyle_v3.py — 严格按杭电毕设论文格式规范重刷 v1 的所有段落。

格式规范来源：用户提供的《毕业设计（论文）参考格式》11 条。
所有尺寸均以 OOXML twips (1/20 pt) 为单位，字号以半磅 (half-point) 为单位。
"""
import shutil, zipfile, os
from pathlib import Path
from copy import deepcopy
from lxml import etree

WORKDIR = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/docx_build')
V1_PATH = WORKDIR / '论文_filled.docx'
TPL_PATH = Path('/mnt/d/Research/Hopfion/09_paper_thesis_talks/bishe/2026届毕业设计模板及相关规定、要求/'
                '毕业设计(论文)材料模板/论文.docx')
OUT_PATH = WORKDIR / '论文_filled_v3.docx'

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

# ═══════════════════════════════════════════════════════════
#  字号表 (w:sz / w:szCs 用半磅)
# ═══════════════════════════════════════════════════════════
SZ_SAN_HAO    = '32'   # 三号 16pt
SZ_XIAO_SAN  = '30'   # 小三 15pt
SZ_SI_HAO    = '28'   # 四号 14pt
SZ_XIAO_SI   = '24'   # 小四 12pt
SZ_WU_HAO    = '21'   # 五号 10.5pt

# 行距 20 磅 = 400 twips, exact
LINE_20PT = '400'

# 段前/段后 2 行 (用 beforeLines/afterLines, 单位 1/100 行)
BEFORE_2LINE = '200'
AFTER_2LINE  = '200'

# 页面 (twips; 1cm = 567 twips)
PAGE_W = '11906'   # A4 width
PAGE_H = '16838'   # A4 height
M_TOP    = '1701'  # 3cm
M_BOTTOM = '1134'  # 2cm
M_LEFT   = '1701'  # 3cm
M_RIGHT  = '1134'  # 2cm
M_GUTTER = '567'   # 装订线 1cm
M_HEADER = '1134'  # 页眉距边界 2cm
M_FOOTER = '567'   # 页脚距边界 1cm


def unpack(p, d):
    if d.exists(): shutil.rmtree(d)
    d.mkdir(parents=True)
    with zipfile.ZipFile(p) as z: z.extractall(d)

def repack(s, o):
    if o.exists(): o.unlink()
    with zipfile.ZipFile(o, 'w', zipfile.ZIP_DEFLATED) as z:
        for root, _, files in os.walk(s):
            for f in files:
                full = Path(root) / f
                z.write(full, full.relative_to(s))

def get_style(p):
    ppr = p.find(f'{W}pPr')
    if ppr is not None:
        ps = ppr.find(f'{W}pStyle')
        if ps is not None: return ps.get(f'{W}val')
    return None

def get_text(p):
    return ''.join(t.text or '' for t in p.iter(f'{W}t'))

def has_drawing(p):
    ns = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    return bool(p.findall(f'.//{ns}inline') or p.findall(f'.//{ns}anchor'))


# ═══════════════════════════════════════════════════════════
#  pPr / rPr 工厂函数
# ═══════════════════════════════════════════════════════════

def _ppr(*children):
    ppr = etree.Element(f'{W}pPr')
    for c in children:
        ppr.append(c)
    return ppr

def _rpr_elem(**kw):
    """Build an <w:rPr> from keyword specs."""
    rpr = etree.Element(f'{W}rPr')
    if 'fonts' in kw:
        rf = etree.SubElement(rpr, f'{W}rFonts')
        for k, v in kw['fonts'].items():
            rf.set(f'{W}{k}', v)
    if kw.get('bold'):
        etree.SubElement(rpr, f'{W}b')
        etree.SubElement(rpr, f'{W}bCs')
    if 'sz' in kw:
        s = etree.SubElement(rpr, f'{W}sz')
        s.set(f'{W}val', kw['sz'])
        s2 = etree.SubElement(rpr, f'{W}szCs')
        s2.set(f'{W}val', kw.get('szCs', kw['sz']))
    return rpr

# ── 规范第 3 条: "摘要"/"目录"/"致谢"/"参考文献"/"附录" 标题 ──
def ppr_section_title_cn():
    """黑体三号, 居中, 单倍行距, 段前2行段后2行"""
    ppr = etree.Element(f'{W}pPr')
    jc = etree.SubElement(ppr, f'{W}jc'); jc.set(f'{W}val', 'center')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}beforeLines', BEFORE_2LINE)
    spc.set(f'{W}afterLines', AFTER_2LINE)
    spc.set(f'{W}line', '240'); spc.set(f'{W}lineRule', 'auto')  # 单倍
    rpr = _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SAN_HAO)
    ppr.append(rpr)
    return ppr

def rpr_section_title_cn():
    return _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SAN_HAO)

# ── 规范第 3 条: "ABSTRACT" 标题 ──
def ppr_section_title_en():
    """Times New Roman 加黑三号, 居中, 单倍行距, 段前2行段后2行"""
    ppr = etree.Element(f'{W}pPr')
    jc = etree.SubElement(ppr, f'{W}jc'); jc.set(f'{W}val', 'center')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}beforeLines', BEFORE_2LINE)
    spc.set(f'{W}afterLines', AFTER_2LINE)
    spc.set(f'{W}line', '240'); spc.set(f'{W}lineRule', 'auto')
    rpr = _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman'},
                    bold=True, sz=SZ_SAN_HAO)
    ppr.append(rpr)
    return ppr

def rpr_section_title_en():
    return _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman'},
                     bold=True, sz=SZ_SAN_HAO)

# ── 规范第 7 条: 第一级标题 (章标题) ──
def ppr_heading1():
    """黑体三号, 居中, 单倍行距, 段前2行段后2行, 前分页"""
    ppr = etree.Element(f'{W}pPr')
    ps = etree.SubElement(ppr, f'{W}pStyle'); ps.set(f'{W}val', '1')
    jc = etree.SubElement(ppr, f'{W}jc'); jc.set(f'{W}val', 'center')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}beforeLines', BEFORE_2LINE)
    spc.set(f'{W}afterLines', AFTER_2LINE)
    spc.set(f'{W}line', '240'); spc.set(f'{W}lineRule', 'auto')
    etree.SubElement(ppr, f'{W}pageBreakBefore')
    rpr = _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SAN_HAO)
    ppr.append(rpr)
    return ppr

def rpr_heading1():
    return _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SAN_HAO)

# ── 规范第 7 条: 第二级标题 ──
def ppr_heading2():
    """黑体四号, 靠左顶格 (1.1 序次), 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    ps = etree.SubElement(ppr, f'{W}pStyle'); ps.set(f'{W}val', '2')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}left', '0'); ind.set(f'{W}firstLine', '0')
    rpr = _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SI_HAO)
    ppr.append(rpr)
    return ppr

def rpr_heading2():
    return _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_SI_HAO)

# ── 规范第 7 条: 第三级标题 ──
def ppr_heading3():
    """黑体小四号, 靠左顶格 (1.1.1 序次), 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    ps = etree.SubElement(ppr, f'{W}pStyle'); ps.set(f'{W}val', '3')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}left', '0'); ind.set(f'{W}firstLine', '0')
    rpr = _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_XIAO_SI)
    ppr.append(rpr)
    return ppr

def rpr_heading3():
    return _rpr_elem(fonts={'eastAsia': '黑体'}, sz=SZ_XIAO_SI)

# ── 规范第 7 条: 正文内容 ──
def ppr_body():
    """宋体小四/TNR小四, 首行缩进2字符, 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}firstLine', '480')
    ind.set(f'{W}firstLineChars', '200')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    rpr = _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman',
                           'eastAsia': '宋体'}, sz=SZ_XIAO_SI)
    ppr.append(rpr)
    return ppr

def rpr_body():
    return _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman',
                            'eastAsia': '宋体'}, sz=SZ_XIAO_SI)

# ── 规范第 5 条: 中文摘要内容 ──
def ppr_cn_abstract():
    """宋体小四, 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}firstLine', '480')
    ind.set(f'{W}firstLineChars', '200')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    rpr = _rpr_elem(fonts={'eastAsia': '宋体'}, sz=SZ_XIAO_SI)
    ppr.append(rpr)
    return ppr

def rpr_cn_abstract():
    return _rpr_elem(fonts={'eastAsia': '宋体'}, sz=SZ_XIAO_SI)

# ── 规范第 5 条: 中文关键词 ──
def ppr_cn_keywords():
    """"关键词"黑体小四, 内容宋体小四, 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    return ppr

# ── 规范第 5 条: 英文摘要内容 ──
def ppr_en_abstract():
    """TNR 小四, 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}firstLine', '480')
    ind.set(f'{W}firstLineChars', '200')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    rpr = _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman'},
                    sz=SZ_XIAO_SI)
    ppr.append(rpr)
    return ppr

def rpr_en_abstract():
    return _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman'},
                     sz=SZ_XIAO_SI)

# ── 规范第 8/9 条: 图注/表注 ──
def ppr_caption():
    """宋体五号, 居中"""
    ppr = etree.Element(f'{W}pPr')
    jc = etree.SubElement(ppr, f'{W}jc'); jc.set(f'{W}val', 'center')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    rpr = _rpr_elem(fonts={'eastAsia': '宋体'}, sz=SZ_WU_HAO)
    ppr.append(rpr)
    return ppr

def rpr_caption():
    return _rpr_elem(fonts={'eastAsia': '宋体'}, sz=SZ_WU_HAO)

# ── 规范第 11 条: 参考文献 ──
def ppr_reference():
    """宋体/TNR 小四, 居左顶格, 行距20磅"""
    ppr = etree.Element(f'{W}pPr')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    ind = etree.SubElement(ppr, f'{W}ind')
    ind.set(f'{W}left', '0'); ind.set(f'{W}firstLine', '0')
    rpr = _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman',
                           'eastAsia': '宋体'}, sz=SZ_XIAO_SI)
    ppr.append(rpr)
    return ppr

def rpr_reference():
    return _rpr_elem(fonts={'ascii': 'Times New Roman', 'hAnsi': 'Times New Roman',
                            'eastAsia': '宋体'}, sz=SZ_XIAO_SI)

# ── 图片所在段 (居中, 无缩进) ──
def ppr_figure():
    ppr = etree.Element(f'{W}pPr')
    jc = etree.SubElement(ppr, f'{W}jc'); jc.set(f'{W}val', 'center')
    spc = etree.SubElement(ppr, f'{W}spacing')
    spc.set(f'{W}line', LINE_20PT); spc.set(f'{W}lineRule', 'exact')
    return ppr


# ═══════════════════════════════════════════════════════════
#  段落格式替换核心
# ═══════════════════════════════════════════════════════════

def restyle_ppr(p, new_ppr):
    old = p.find(f'{W}pPr')
    sect = None
    if old is not None:
        sect = old.find(f'{W}sectPr')
        if sect is not None: sect = deepcopy(sect)
        p.remove(old)
    fresh = deepcopy(new_ppr)
    if sect is not None:
        fresh.append(sect)
    p.insert(0, fresh)

def restyle_runs(p, new_rpr):
    KEEP = {f'{W}b', f'{W}bCs', f'{W}i', f'{W}iCs',
            f'{W}vertAlign', f'{W}u', f'{W}strike', f'{W}smallCaps'}
    for r in p.findall(f'{W}r'):
        if r.find(f'{W}drawing') is not None or r.find(f'{W}pict') is not None:
            continue
        old = r.find(f'{W}rPr')
        kept = {}
        if old is not None:
            for child in old:
                if child.tag in KEEP:
                    kept[child.tag] = deepcopy(child)
            r.remove(old)
        fresh = deepcopy(new_rpr)
        for tag in KEEP:
            for e in fresh.findall(tag):
                fresh.remove(e)
        for tag, elem in kept.items():
            fresh.append(elem)
        r.insert(0, fresh)


# ═══════════════════════════════════════════════════════════
#  页面设置 (sectPr)
# ═══════════════════════════════════════════════════════════

def fix_sect_pr(sect):
    """Fix page margins in a <w:sectPr>."""
    pg = sect.find(f'{W}pgSz')
    if pg is None:
        pg = etree.SubElement(sect, f'{W}pgSz')
    pg.set(f'{W}w', PAGE_W); pg.set(f'{W}h', PAGE_H)

    mg = sect.find(f'{W}pgMar')
    if mg is None:
        mg = etree.SubElement(sect, f'{W}pgMar')
    mg.set(f'{W}top', M_TOP)
    mg.set(f'{W}bottom', M_BOTTOM)
    mg.set(f'{W}left', M_LEFT)
    mg.set(f'{W}right', M_RIGHT)
    mg.set(f'{W}gutter', M_GUTTER)
    mg.set(f'{W}header', M_HEADER)
    mg.set(f'{W}footer', M_FOOTER)


# ═══════════════════════════════════════════════════════════
#  Main
# ═══════════════════════════════════════════════════════════

def main():
    v1_dir = WORKDIR / '_v1_v3'
    tpl_dir = WORKDIR / '_tpl_v3'

    print("Step 1: Unpacking...")
    unpack(V1_PATH, v1_dir)
    unpack(TPL_PATH, tpl_dir)

    # Copy template's styles, headers, footers
    print("Step 2: Copying template styles/headers/footers...")
    for f in ['styles.xml']:
        src = tpl_dir / 'word' / f
        if src.exists():
            shutil.copy2(src, v1_dir / 'word' / f)
    for src in (tpl_dir / 'word').glob('header*.xml'):
        shutil.copy2(src, v1_dir / 'word' / src.name)
    for src in (tpl_dir / 'word').glob('footer*.xml'):
        shutil.copy2(src, v1_dir / 'word' / src.name)

    # Parse document
    v1_doc = etree.parse(str(v1_dir / 'word' / 'document.xml'))
    v1_body = v1_doc.getroot().find(f'{W}body')

    # Step 3: Fix all sectPr (page margins)
    print("Step 3: Fixing page margins...")
    for sect in v1_body.findall(f'.//{W}sectPr'):
        fix_sect_pr(sect)
    print("  All sectPr updated")

    # Step 4: Restyle all paragraphs
    print("Step 4: Restyling paragraphs...")
    paras = v1_body.findall(f'{W}p')
    context = 'cover'
    stats = {}

    # Section title lines to detect
    SECTION_TITLES_CN = {'摘要', '摘    要', '目录', '目    录', '致谢', '致    谢',
                         '参考文献', '附录'}

    for i, p in enumerate(paras):
        style = get_style(p)
        txt = get_text(p).strip()

        # ── Context transitions ──
        # Cover ends at 承诺页 "年 月 日" line; skip everything before
        if context == 'cover':
            if '年' in txt and '月' in txt and '日' in txt:
                context = 'post_cover'
            continue
        if context == 'post_cover' and not txt:
            continue  # skip blank lines between cover and abstract

        if txt in ('摘要', '摘    要'):
            context = 'cn_abstract'
            restyle_ppr(p, ppr_section_title_cn())
            restyle_runs(p, rpr_section_title_cn())
            stats['section_title'] = stats.get('section_title', 0) + 1
            continue

        if txt == 'ABSTRACT':
            context = 'en_abstract'
            restyle_ppr(p, ppr_section_title_en())
            restyle_runs(p, rpr_section_title_en())
            stats['section_title'] = stats.get('section_title', 0) + 1
            continue

        if '目' in txt and '录' in txt and i < 80:
            context = 'toc'
            restyle_ppr(p, ppr_section_title_cn())
            restyle_runs(p, rpr_section_title_cn())
            stats['section_title'] = stats.get('section_title', 0) + 1
            continue

        if style in ('10', '22'):  # TOC entry styles
            context = 'toc'
            continue

        # First heading after abstract/toc → body
        if style in ('1',) and context in ('toc', 'en_abstract', 'cn_abstract', 'post_cover'):
            context = 'body'

        if context == 'en_abstract' and style in ('1', '2', '3'):
            context = 'body'

        if style in ('1',) and '致谢' in txt:
            context = 'ack_title'
        elif style in ('1',) and '参考文献' in txt:
            context = 'ref_title'
        elif style in ('1',) and '附录' in txt:
            context = 'appendix'

        # Detect refs by style or content pattern
        # (pandoc uses 'Bibliography' style, or refs appear after ack)
        if style == 'Bibliography':
            context = 'refs'
        elif context == 'ack' and txt:
            import re
            if re.match(r'^[A-Z\u4e00-\u9fff].*\d{4}[.)\"]', txt) and len(txt) > 40:
                context = 'refs'

        if context in ('cover', 'post_cover', 'toc', 'appendix'):
            continue

        # ── Apply formatting by context + role ──

        # Special section title headings
        if context == 'ack_title':
            restyle_ppr(p, ppr_section_title_cn())
            restyle_runs(p, rpr_section_title_cn())
            stats['section_title'] = stats.get('section_title', 0) + 1
            context = 'ack'
            continue

        if context == 'ref_title':
            restyle_ppr(p, ppr_section_title_cn())
            restyle_runs(p, rpr_section_title_cn())
            stats['section_title'] = stats.get('section_title', 0) + 1
            context = 'refs'
            continue

        # CN Abstract content
        if context == 'cn_abstract':
            if '关键词' in txt:
                restyle_ppr(p, ppr_cn_keywords())
                # Don't restyle runs — "关键词" should be bold 黑体, rest 宋体
                stats['cn_keywords'] = 1
                continue
            restyle_ppr(p, ppr_cn_abstract())
            restyle_runs(p, rpr_cn_abstract())
            stats['cn_abstract'] = stats.get('cn_abstract', 0) + 1
            continue

        # EN Abstract content
        if context == 'en_abstract':
            if 'key word' in txt.lower() or 'keywords' in txt.lower():
                restyle_ppr(p, ppr_cn_keywords())
                stats['en_keywords'] = 1
                continue
            restyle_ppr(p, ppr_en_abstract())
            restyle_runs(p, rpr_en_abstract())
            stats['en_abstract'] = stats.get('en_abstract', 0) + 1
            continue

        # References
        if context == 'refs':
            restyle_ppr(p, ppr_reference())
            restyle_runs(p, rpr_reference())
            stats['reference'] = stats.get('reference', 0) + 1
            continue

        # Acknowledgment (same as body text)
        if context == 'ack':
            restyle_ppr(p, ppr_body())
            restyle_runs(p, rpr_body())
            stats['ack'] = stats.get('ack', 0) + 1
            continue

        # Body content
        if context == 'body':
            if style in ('1',):
                restyle_ppr(p, ppr_heading1())
                restyle_runs(p, rpr_heading1())
                stats['heading1'] = stats.get('heading1', 0) + 1
            elif style in ('2',):
                restyle_ppr(p, ppr_heading2())
                restyle_runs(p, rpr_heading2())
                stats['heading2'] = stats.get('heading2', 0) + 1
            elif style in ('3', 'Heading3', 'Heading4'):
                restyle_ppr(p, ppr_heading3())
                restyle_runs(p, rpr_heading3())
                stats['heading3'] = stats.get('heading3', 0) + 1
            elif has_drawing(p):
                restyle_ppr(p, ppr_figure())
                stats['figure'] = stats.get('figure', 0) + 1
            elif txt and (txt.startswith('图') or txt.startswith('表')) and len(txt) < 120:
                restyle_ppr(p, ppr_caption())
                restyle_runs(p, rpr_caption())
                stats['caption'] = stats.get('caption', 0) + 1
            else:
                restyle_ppr(p, ppr_body())
                restyle_runs(p, rpr_body())
                stats['body'] = stats.get('body', 0) + 1

    print(f"  Stats: {stats}")

    # Step 5: Write and repack
    print("\nStep 5: Repacking...")
    etree.ElementTree(v1_doc.getroot()).write(
        str(v1_dir / 'word' / 'document.xml'),
        xml_declaration=True, encoding='UTF-8', standalone=True)
    repack(v1_dir, OUT_PATH)

    sz = OUT_PATH.stat().st_size / 1024 / 1024
    print(f"\nDONE: {OUT_PATH} ({sz:.1f} MB)")

    # Copy to target
    targets = [
        TPL_PATH.parent / '论文_filled_v3.docx',
    ]
    for t in targets:
        try:
            shutil.copy2(OUT_PATH, t)
            print(f"Copied to: {t}")
        except PermissionError:
            print(f"  WARN: {t} locked")

    # Verify
    from docx import Document
    doc = Document(str(OUT_PATH))
    print(f"Verify: {len(doc.paragraphs)} paras, {len(doc.tables)} tables")

    shutil.rmtree(v1_dir, ignore_errors=True)
    shutil.rmtree(tpl_dir, ignore_errors=True)


if __name__ == '__main__':
    main()
